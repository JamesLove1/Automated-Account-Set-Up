#%%
from docx import Document
from docx.shared import Pt
import os
#%%
def letterOne(company, position,doc, output):
    print(os.getcwd())
    
    #file_path = "./v1 Cover Letter Tech - James Love Cover Letter.docx"
    
    document = Document(doc)

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)

    paragraph_1 = document.paragraphs[1]
    paragraph_1.style = document.styles['Normal']
    paragraph_1.text = "I am writing to express my strong interest in the "+position+" position at "+company+". Having self-taught myself various coding languages – which gives me a great advantage when it comes to developing scripts and troubleshooting issues –, I have now applied to computer science courses at university. As a passionate and driven individual, I am confident that my enthusiasm for technology and my communication skills make me an ideal candidate for this role."

    paragraph_4 = document.paragraphs[4]
    paragraph_4.style = document.styles['Normal']
    paragraph_4.text = "Overall, I am eager to join the "+company+" team and I am confident that I can make a meaningful contribution. I have the required skills, knowledge, and enthusiasm to make a positive impact on the team. I am looking forward to the opportunity to discuss my application further."

    document.save(output+"/"+company+" - James Love Cover Letter.docx")
    #document.save("output within framwork folder.docx")

# %%
